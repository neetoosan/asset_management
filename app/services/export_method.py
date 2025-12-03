import csv
import os
from datetime import datetime
from typing import Any
try:
    import docx.enum.text
except:
    pass

try:
    import openpyxl
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def _is_grouped(data: Any) -> bool:
    return isinstance(data, dict)


def export_csv(data: Any, file_path: str, report_title: str = "Report") -> None:
    """Export data to CSV. Supports grouped data (dict of lists) or list of dicts."""
    try:
        if not data:
            raise ValueError("No data provided for export")
        if not file_path:
            raise ValueError("No file path specified")
    except ValueError as ve:
        print(f"CSV Export Error: {ve}")
        raise
    
    os.makedirs(os.path.dirname(file_path) or '.', exist_ok=True)
    with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow([f"{report_title} - Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}"])
        writer.writerow([])

        if _is_grouped(data):
            for group, rows in data.items():
                writer.writerow([f"=== {group} ==="])
                if rows:
                    if isinstance(rows[0], dict):
                        writer.writerow(list(rows[0].keys()))
                        for r in rows:
                            writer.writerow([r.get(k, '') for k in rows[0].keys()])
                    else:
                        for r in rows:
                            writer.writerow(r)
                writer.writerow([])
        else:
            if data:
                if isinstance(data[0], dict):
                    fieldnames = list(data[0].keys())
                    writer.writerow(fieldnames)
                    for row in data:
                        writer.writerow([row.get(k, '') for k in fieldnames])
                else:
                    for row in data:
                        writer.writerow(row)


def export_xlsx(data: Any, file_path: str, report_title: str = "Report") -> None:
    """Export data to XLSX with professional formatting, freeze panes, and auto-fit columns."""
    try:
        if not data:
            raise ValueError("No data provided for Excel export")
        if not file_path:
            raise ValueError("No file path specified for Excel export")
    except ValueError as ve:
        print(f"Excel Export Error: {ve}")
        raise
    
    if not OPENPYXL_AVAILABLE:
        # fallback
        print("Warning: openpyxl not available, falling back to CSV")
        fallback = file_path.replace('.xlsx', '.csv')
        export_csv(data, fallback, report_title)
        return

    wb = Workbook()
    ws = wb.active
    ws.title = report_title[:31]

    # Professional styles
    header_font = Font(bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill(start_color='2E86C1', end_color='2E86C1', fill_type='solid')
    title_font = Font(bold=True, color='2E86C1', size=14)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
    right_align = Alignment(horizontal='right', vertical='center')
    thin = Side(border_style='thin', color='D3D3D3')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    currency_format = '#,##0.00'  # Simplified format
    date_format = 'yyyy-mm-dd'

    row_cursor = 1
    # Title row
    max_cols = 10
    ws.merge_cells(start_row=row_cursor, start_column=1, end_row=row_cursor, end_column=max_cols)
    title_cell = ws.cell(row=row_cursor, column=1, value=f"{report_title}")
    title_cell.font = title_font
    ws.row_dimensions[row_cursor].height = 24
    row_cursor += 1
    
    # Generated date row
    ws.merge_cells(start_row=row_cursor, start_column=1, end_row=row_cursor, end_column=max_cols)
    date_cell = ws.cell(row=row_cursor, column=1, value=f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_cell.font = Font(size=10, color='666666')
    ws.row_dimensions[row_cursor].height = 16
    row_cursor += 2

    if _is_grouped(data):
        for group, rows in data.items():
            ws.cell(row=row_cursor, column=1, value=f"{group}")
            row_cursor += 1
            if not rows:
                row_cursor += 1
                continue
            if isinstance(rows[0], dict):
                headers = list(rows[0].keys())
                for c, h in enumerate(headers, start=1):
                    cell = ws.cell(row=row_cursor, column=c, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center
                    cell.border = border
                row_cursor += 1
                for r in rows:
                    for c, h in enumerate(headers, start=1):
                        cell = ws.cell(row=row_cursor, column=c, value=r.get(h, ''))
                        cell.border = border
                    row_cursor += 1
            else:
                for r in rows:
                    ws.cell(row=row_cursor, column=1, value=str(r))
                    row_cursor += 1
            row_cursor += 1
    else:
        if not data:
            pass
        else:
            if isinstance(data[0], dict):
                headers = list(data[0].keys())
                for c, h in enumerate(headers, start=1):
                    cell = ws.cell(row=row_cursor, column=c, value=h)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center
                    cell.border = border
                row_cursor += 1
                for r in data:
                    for c, h in enumerate(headers, start=1):
                        cell = ws.cell(row=row_cursor, column=c, value=r.get(h, ''))
                        cell.border = border
                    row_cursor += 1
            else:
                for r in data:
                    ws.cell(row=row_cursor, column=1, value=str(r))
                    row_cursor += 1

    # Freeze header rows (set to first data row)
    try:
        ws.freeze_panes = 'A4'  # Freeze top 3 rows
    except Exception:
        pass  # Skip if freeze fails
    
    # Auto-fit columns with intelligent sizing
    try:
        max_col = ws.max_column
        for col_idx in range(1, max_col + 1):
            col_letter = get_column_letter(col_idx)
            max_length = 0
            is_numeric = False
            
            for row_idx, row in enumerate(ws.iter_rows(min_col=col_idx, max_col=col_idx)):
                cell = row[0]
                try:
                    val = cell.value
                    if val is not None:
                        val_str = str(val)
                        if len(val_str) > max_length:
                            max_length = len(val_str)
                        # Check if numeric for formatting
                        if isinstance(val, (int, float)) and row_idx > 0:
                            is_numeric = True
                            if isinstance(val, float):
                                cell.number_format = '0.00'  # Decimal format
                            else:
                                cell.number_format = '0'  # Integer format
                except Exception:
                    continue
            
            # Set column width with minimum
            adjusted_width = max(max_length + 2, 12)
            ws.column_dimensions[col_letter].width = min(adjusted_width, 50)  # Cap at 50
    except Exception:
        pass

    wb.save(file_path)


def export_pdf(data: Any, file_path: str, report_title: str = "Report") -> None:
    """Export data to PDF in A4 landscape format with professional styling."""
    try:
        if not data:
            raise ValueError("No data provided for PDF export")
        if not file_path:
            raise ValueError("No file path specified for PDF export")
        if not report_title or not isinstance(report_title, str):
            report_title = "Report"
    except ValueError as ve:
        print(f"PDF Export Error: {ve}")
        raise
    
    os.makedirs(os.path.dirname(file_path) or '.', exist_ok=True)
    if not REPORTLAB_AVAILABLE:
        print("Warning: reportlab not available, generating text report instead")
        # fallback: write plain text
        txt_path = file_path.replace('.pdf', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"{report_title}\nGenerated on {datetime.now().isoformat()}\n\n")
            if _is_grouped(data):
                for group, rows in data.items():
                    f.write(f"== {group} ==\n")
                    for r in rows:
                        f.write(str(r) + '\n')
                    f.write('\n')
            else:
                for r in data:
                    f.write(str(r) + '\n')
        return

    # Use reportlab to generate professional PDF in A4 landscape
    pagesize = landscape(A4)  # A4 landscape for wide tables
    doc = SimpleDocTemplate(file_path, pagesize=pagesize, topMargin=0.5*inch, bottomMargin=0.5*inch,
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    # Create custom title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#2E86C1'),
        spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=12
    )
    
    elems = []
    elems.append(Paragraph(report_title, title_style))
    elems.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
    elems.append(Spacer(1, 8))

    if _is_grouped(data):
        for group, rows in data.items():
            if rows:  # Only add section if it has data
                elems.append(Paragraph(str(group), styles['Heading3']))
                if isinstance(rows[0], dict):
                    headers = list(rows[0].keys())
                    # Calculate column widths based on page width
                    page_width = pagesize[0] - 1*inch
                    col_width = page_width / len(headers)
                    col_widths = [col_width] * len(headers)
                    
                    table_data = [headers]
                    for r in rows:
                        row = []
                        for h in headers:
                            val = r.get(h, '')
                            # Format numbers and currencies
                            if isinstance(val, (int, float)):
                                if isinstance(val, float) and val != int(val):
                                    val = f"{val:,.2f}"  # Currency format
                                else:
                                    val = f"{val:,}"  # Number format
                            row.append(str(val))
                        table_data.append(row)
                else:
                    col_widths = [pagesize[0] - 1*inch]
                    table_data = [[str(r)] for r in rows]

                t = Table(table_data, colWidths=col_widths, hAlign='CENTER')
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86C1')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
                    ('TOPPADDING', (0, 1), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                elems.append(t)
                elems.append(Spacer(1, 12))
    else:
        if data:
            if isinstance(data[0], dict):
                headers = list(data[0].keys())
                # Calculate column widths
                page_width = pagesize[0] - 1*inch
                col_width = page_width / len(headers)
                col_widths = [col_width] * len(headers)
                
                table_data = [headers]
                for r in data:
                    row = []
                    for h in headers:
                        val = r.get(h, '')
                        # Format numbers and currencies
                        if isinstance(val, (int, float)):
                            if isinstance(val, float) and val != int(val):
                                val = f"{val:,.2f}"  # Currency format
                            else:
                                val = f"{val:,}"  # Number format
                        row.append(str(val))
                    table_data.append(row)
            else:
                col_widths = [pagesize[0] - 1*inch]
                table_data = [[str(r)] for r in data]

            t = Table(table_data, colWidths=col_widths, hAlign='CENTER')
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E86C1')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
                ('TOPPADDING', (0, 1), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
            ]))
            elems.append(t)
        else:
            elems.append(Paragraph('No data to display', styles['Normal']))

    doc.build(elems)


def export_docx(data: Any, file_path: str, report_title: str = "Report") -> None:
    """Export data to DOCX with professional formatting, tables, and headers."""
    try:
        if not data:
            raise ValueError("No data provided for Word export")
        if not file_path:
            raise ValueError("No file path specified for Word export")
        if not report_title or not isinstance(report_title, str):
            report_title = "Report"
    except ValueError as ve:
        print(f"Word Export Error: {ve}")
        raise
    
    os.makedirs(os.path.dirname(file_path) or '.', exist_ok=True)
    if not DOCX_AVAILABLE:
        print("Warning: python-docx not available, generating text report instead")
        txt_path = file_path.replace('.docx', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"{report_title}\nGenerated on {datetime.now().isoformat()}\n\n")
            if _is_grouped(data):
                for group, rows in data.items():
                    f.write(f"== {group} ==\n")
                    for r in rows:
                        f.write(str(r) + '\n')
                    f.write('\n')
            else:
                for r in data:
                    f.write(str(r) + '\n')
        return

    doc = docx.Document()
    # Add professional header
    title = doc.add_heading(report_title, level=0)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    timestamp_format = timestamp.runs[0].font
    timestamp_format.italic = True
    timestamp_format.size = Pt(10)
    
    doc.add_paragraph()  # Blank line for spacing

    if _is_grouped(data):
        for group, rows in data.items():
            doc.add_heading(str(group), level=2)
            if not rows:
                continue
            if isinstance(rows[0], dict):
                headers = list(rows[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'  # Professional table style
                
                # Format header
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(hdr_cells):
                    h.text = str(headers[i])
                    # Bold header text
                    for paragraph in h.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows
                for r in rows:
                    row_cells = table.add_row().cells
                    for i, h in enumerate(headers):
                        val = r.get(h, '')
                        # Format numeric values
                        if isinstance(val, (int, float)):
                            if isinstance(val, float):
                                val = f"{val:,.2f}"
                            else:
                                val = f"{val:,}"
                        row_cells[i].text = str(val)
            else:
                for r in rows:
                    doc.add_paragraph(str(r))
            doc.add_paragraph()  # Spacing between sections
    else:
        if not data:
            doc.add_paragraph('No data to display')
        else:
            if isinstance(data[0], dict):
                headers = list(data[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Format header
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(hdr_cells):
                    h.text = str(headers[i])
                    for paragraph in h.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows
                for r in data:
                    row_cells = table.add_row().cells
                    for i, h in enumerate(headers):
                        val = r.get(h, '')
                        if isinstance(val, (int, float)):
                            if isinstance(val, float):
                                val = f"{val:,.2f}"
                            else:
                                val = f"{val:,}"
                        row_cells[i].text = str(val)
            else:
                for r in data:
                    doc.add_paragraph(str(r))

    doc.save(file_path)
